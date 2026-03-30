import os
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Force CPU
os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
os.environ['DISABLE_MODEL_SOURCE_CHECK'] = 'True'
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'

import tempfile
import uuid
import threading
import time
import shutil
import subprocess
import atexit
import gc
import ctypes
from pathlib import Path
from flask import Flask, request, jsonify, send_file, Response, after_this_request
import fitz  # PyMuPDF
import numpy as np
import re

# ── Config ───────────────────────────────────────────────────────────────────
MAX_UPLOAD_MB    = int(os.environ.get('MAX_UPLOAD_MB', 200))
CLEANUP_INTERVAL = int(os.environ.get('CLEANUP_INTERVAL', 3600))
MAX_FILE_AGE     = int(os.environ.get('MAX_FILE_AGE', 3600))
JOB_MAX_AGE      = int(os.environ.get('JOB_MAX_AGE', 7200))
OCR_DPI          = int(os.environ.get('OCR_DPI', 200))
OCR_BATCH_SIZE   = int(os.environ.get('OCR_BATCH_SIZE', 2))
CPU_THREADS      = int(os.environ.get('CPU_THREADS', 8))
MIN_CONFIDENCE   = float(os.environ.get('MIN_CONFIDENCE', 0.5))
# PP-OCRv5 推理增强开关（可通过 docker-compose 环境变量动态控制）
USE_DOC_ORIENTATION = os.environ.get('USE_DOC_ORIENTATION', 'true').lower() == 'true'
USE_DOC_UNWARPING   = os.environ.get('USE_DOC_UNWARPING',   'true').lower() == 'true'
USE_TEXTLINE_ORI    = os.environ.get('USE_TEXTLINE_ORI',    'false').lower() == 'true' 
# 文字层判断阈值：提取到的字符数超过此值才认为有文字层
TEXT_LAYER_MIN_CHARS = int(os.environ.get('TEXT_LAYER_MIN_CHARS', 20))
PDF_PPI          = 72

ALLOWED_EXTENSIONS = {'.pdf', '.jpg', '.jpeg', '.png'}

# ── Flask ─────────────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder='static', static_url_path='/static')
app.config['MAX_CONTENT_LENGTH'] = MAX_UPLOAD_MB * 1024 * 1024
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(32).hex())

# ── Storage ───────────────────────────────────────────────────────────────────
OUTPUT_DIR = Path(tempfile.gettempdir()) / 'pdf_ocr_output'
UPLOAD_DIR = Path(tempfile.gettempdir()) / 'pdf_ocr_uploads'
OUTPUT_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

# ── Job / Progress state ──────────────────────────────────────────────────────
jobs: dict          = {}
jobs_lock           = threading.Lock()
progress_data: dict = {}
progress_lock       = threading.Lock()
cancel_flags: dict  = {}
cancel_lock         = threading.Lock()

# Serialise OCR calls — PaddleOCR is not thread-safe
ocr_lock = threading.Lock()

# ── Model lifecycle ───────────────────────────────────────────────────────────
_ocr               = None
_model_lock        = threading.Lock()
_ocr_last_activity = 0.0
MODEL_IDLE_TIMEOUT = int(os.environ.get('MODEL_IDLE_TIMEOUT', 3600))

def full_memory_release():
    global _ocr
    try:
        if _ocr is not None:
            del _ocr
            _ocr = None
        gc.collect()
        # Linux/NAS系统级回收内存
        libc = ctypes.CDLL("libc.so.6")
        libc.malloc_trim(0)
    except:
        pass

def get_ocr():
    """Return loaded OCR instance, raise if not loaded."""
    global _ocr_last_activity
    if _ocr is None:
        raise RuntimeError("OCR 引擎未启动，请先在页面开启引擎")
    _ocr_last_activity = time.time()
    return _ocr


def load_model(use_doc_ori=USE_DOC_ORIENTATION, 
               use_unwarp=USE_DOC_UNWARPING, 
               use_text_ori=USE_TEXTLINE_ORI):
    """
    Manually load PaddleOCR model.
    根据前端传入的参数动态决定加载哪些子模型。
    """
    global _ocr, _ocr_last_activity
    with _model_lock:
        if _ocr is not None:
            return 'already_loaded'
        try:
            logger.info(
                f"Model: Loading PaddleOCR (PP-OCRv5, CPU)… "
                f"orientation={use_doc_ori} "
                f"unwarping={use_unwarp} "
                f"textline_ori={use_text_ori}"
            )
            # 防止触发 MKLDNN 报错Bug（如果之前已经加了请保留）
            os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
            os.environ['FLAGS_use_onednn'] = '0'
            os.environ['FLAGS_use_mkldnn'] = '0'

            from paddleocr import PaddleOCR
            _ocr = PaddleOCR(
                # 轻量级 mobile 模型，内存友好
                text_detection_model_name="PP-OCRv5_mobile_det",
                text_recognition_model_name="PP-OCRv5_mobile_rec",
                use_doc_orientation_classify=use_doc_ori,
                use_doc_unwarping=use_unwarp,
                use_textline_orientation=use_text_ori,
                device='cpu',
                enable_mkldnn=False,    # 禁用 OneDNN，避免 CPU 算子兼容性报错
            )
            _ocr_last_activity = time.time()
            logger.info("Model: PaddleOCR PP-OCRv5 (Mobile) ready.")
            return 'loaded'
        except Exception:
            _ocr = None
            raise


def unload_model():
    """
    手动卸载 PaddleOCR 模型并强制回收系统内存。
    锁内只做引用断开，耗时的 GC 和 malloc_trim 在锁外执行。
    """
    global _ocr, _ocr_last_activity

    with _model_lock:
        if _ocr is None:
            logger.info("Model: 引擎当前未加载，无需释放。")
            return
        logger.info("Model: 正在执行深度内存释放逻辑...")
        del _ocr
        _ocr = None
        _ocr_last_activity = 0.0

    gc.collect()

    try:
        import paddle
        paddle.device.cuda.empty_cache()
    except Exception:
        pass

    try:
        libc = ctypes.CDLL("libc.so.6")
        res  = libc.malloc_trim(0)
        if res == 1:
            logger.info("System: malloc_trim 成功，内存已归还操作系统。")
        else:
            logger.info("System: malloc_trim 已执行，当前无连续内存块可回收。")
    except Exception as e:
        logger.warning(f"System: malloc_trim 调用失败（非 Linux 环境正常）: {e}")

    logger.info("Model: OCR 引擎已进入待机（壳）状态。")


# ── S2T (simplified → traditional) ───────────────────────────────────────────
try:
    from s2t_dict import S2T_ONE_TO_ONE
    _S2T_TABLE = str.maketrans(S2T_ONE_TO_ONE)
    _HAS_S2T = True
except ImportError:
    _HAS_S2T = False

def fix_ocr_text(text: str) -> str:
    return text


# ── Helpers ───────────────────────────────────────────────────────────────────
def validate_path(filename: str) -> Path:
    resolved = (OUTPUT_DIR / filename).resolve()
    if not resolved.is_relative_to(OUTPUT_DIR.resolve()):
        raise ValueError(f"Invalid path: {filename}")
    return resolved


def sanitize_name(name: str, max_len: int = 180) -> str:
    name = re.sub(r'[\x00-\x1f\x7f]', '', name)
    name = name.replace('"', "'").replace('\\', '_').replace('/', '_')
    name = re.sub(r'[_\s]+', '_', name).strip(' ._')
    return (name[:max_len].rstrip(' ._') or 'download')


def pixmap_to_numpy(pix: fitz.Pixmap) -> np.ndarray:
    if pix.alpha:
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 4)[:, :, :3]
    else:
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
    return img[:, :, ::-1].copy()   # RGB → BGR


def validate_pdf(file) -> bool:
    file.seek(0)
    magic = file.read(5)
    file.seek(0)
    return magic == b'%PDF-'


def is_image_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in {'.jpg', '.jpeg', '.png'}


def load_image_as_numpy(src_path: str) -> np.ndarray:
    """Load JPG/PNG as BGR numpy array via PyMuPDF (no cv2 dependency)."""
    doc = fitz.open(src_path)
    page = doc[0]
    pix  = page.get_pixmap()
    img  = pixmap_to_numpy(pix)

    doc.close()
    return img


# ── Text layer detection ──────────────────────────────────────────────────────
def has_text_layer(src_path: str) -> bool:
    """
    检测 PDF 是否含有可用文字层。
    提取所有页面文字，总字符数超过阈值则认为有文字层。
    避免把水印、页码等少量字符误判为有效文字层。
    """
    try:
        doc = fitz.open(src_path)
        total_chars = 0
        for page in doc:
            text = page.get_text("text")
            total_chars += len(text.strip())
            if total_chars >= TEXT_LAYER_MIN_CHARS:
                doc.close()
                return True
        doc.close()
        return False
    except Exception as e:
        logger.warning(f"has_text_layer check failed: {e}")
        return False


def extract_text_layer(src_path: str) -> str:
    """
    直接从 PDF 文字层提取文字，按 Y 坐标排序拼接。
    用于有文字层的 PDF，完全跳过 OCR。
    """
    doc = fitz.open(src_path)
    all_parts = []

    for page_num, page in enumerate(doc):
        # get_text("blocks") 返回 (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text("blocks")
        if not blocks:
            continue

        # 只取文字块（block_type == 0），过滤图片块（block_type == 1）
        text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]

        # 按 Y 坐标排序（四舍五入到 5px 容差内，再按 X 排序）
        text_blocks.sort(key=lambda b: (round(b[1] / 5) * 5, b[0]))

        page_lines = []
        for b in text_blocks:
            # 块内文字可能含换行，保留原始换行但清理首尾空白
            lines = [ln.strip() for ln in b[4].splitlines() if ln.strip()]
            page_lines.extend(lines)

        if page_lines:
            all_parts.append('\n'.join(page_lines))

    doc.close()
    return '\n\n'.join(all_parts)


# ── Improved line grouping (Y-center clustering) ──────────────────────────────
def _group_lines_v2(items: list) -> list[list]:
    """
    基于 Y 中心点跳变聚类的行分组算法。
    替代旧的逐行比较逻辑，对手机拍摄/梯形畸变图像有更强鲁棒性。

    算法：
    1. 计算每个 OCR 块的 Y 中心点和高度
    2. 按 Y 中心点排序
    3. 动态阈值 = 所有块平均高度 × 0.5
    4. 相邻两块 Y 中心点差值 > 阈值则切割为新行
    5. 同行内按 X 坐标从左到右排序
    """
    if not items:
        return []

    # 计算每个块的 Y 中心点和高度
    def y_center(it):
        ys = [p[1] for p in it['poly']]
        return (min(ys) + max(ys)) / 2.0

    def height(it):
        ys = [p[1] for p in it['poly']]
        return max(ys) - min(ys)

    # 按 Y 中心点排序
    sorted_items = sorted(items, key=y_center)

    # 动态阈值：平均行高 × 0.5，最小为 8px
    heights = [height(it) for it in sorted_items]
    avg_h   = sum(heights) / len(heights) if heights else 20
    threshold = max(8.0, avg_h * 0.5)

    # 聚类
    lines = []
    cur   = [sorted_items[0]]
    for item in sorted_items[1:]:
        prev_cy = y_center(cur[-1])
        cur_cy  = y_center(item)
        if abs(cur_cy - prev_cy) <= threshold:
            cur.append(item)
        else:
            lines.append(cur)
            cur = [item]
    lines.append(cur)

    # 同行内按 X 从左到右排序
    for line in lines:
        line.sort(key=lambda it: min(p[0] for p in it['poly']))

    return lines


# ── Header / footer stripper ─────────────────────────────────────────────────
def _strip_header_footer(items: list, page_h: float,
                         top_ratio: float = 0.08,
                         bottom_ratio: float = 0.08) -> list:
    """
    过滤页眉页脚：
    - 去除 Y 中心点在页面顶部 top_ratio 或底部 bottom_ratio 范围内的文字块
    - 辅助条件：文字较短（< 30字）或匹配页码模式
    只有同时满足「位置在边缘」才过滤，避免误删正文第一行/最后一行。
    """
    if not items or page_h <= 0:
        return items

    import re as _re
    _PAGE_NUM_RE = _re.compile(
        r'^[\s\-\–—·•]*\d+[\s\-\–—·•]*$'   # 纯页码：- 5 - / 5 / · 5 ·
        r'|第\s*\d+\s*[页面]'                # 第5页
        r'|\d+\s*/\s*\d+'                    # 5/10
    )

    top_cut    = page_h * top_ratio
    bottom_cut = page_h * (1 - bottom_ratio)

    def y_center(it):
        ys = [p[1] for p in it['poly']]
        return (min(ys) + max(ys)) / 2.0

    def is_edge(it):
        cy = y_center(it)
        return cy < top_cut or cy > bottom_cut

    def is_short_or_pagenum(it):
        t = it['text'].strip()
        return len(t) < 30 or bool(_PAGE_NUM_RE.search(t))

    filtered = [it for it in items if not (is_edge(it) and is_short_or_pagenum(it))]
    removed  = len(items) - len(filtered)
    if removed:
        import logging as _logging
        _logging.getLogger(__name__).debug(f"strip_header_footer: removed {removed} items")
    return filtered


# ── Progress / cancel ─────────────────────────────────────────────────────────
def update_progress(task_id: str, current: int, total: int, status: str, message: str = ''):
    if status != 'cancelled':
        with progress_lock:
            if progress_data.get(task_id, {}).get('status') == 'cancelled':
                return
        if is_cancelled(task_id):
            return
    now = time.time()
    with jobs_lock:
        job = jobs.get(task_id)
        if job and job.get('status') == 'processing':
            job['updated_at'] = now
    with progress_lock:
        progress_data[task_id] = {
            'current': current,
            'total': total,
            'percent': min(100, int(current / total * 100)) if total else 0,
            'status': status,
            'message': message,
            'updated_at': now,
        }


def is_cancelled(jid: str) -> bool:
    with cancel_lock:
        return cancel_flags.get(jid, False)


def set_cancelled(jid: str, v: bool = True):
    with cancel_lock:
        cancel_flags[jid] = v


# ── Background cleanup ────────────────────────────────────────────────────────
_cleanup_stop = threading.Event()


def _cleanup_loop():
    while not _cleanup_stop.is_set():
        try:
            now = time.time()
            for p in list(OUTPUT_DIR.iterdir()):
                if (now - p.stat().st_mtime) > MAX_FILE_AGE:
                    try:
                        shutil.rmtree(p) if p.is_dir() else p.unlink()
                    except Exception:
                        pass
            with jobs_lock:
                stale = [
                    jid for jid, j in jobs.items()
                    if j.get('status') != 'processing'
                    and now - j.get('updated_at', j.get('created_at', 0)) > JOB_MAX_AGE
                ]
                for jid in stale:
                    del jobs[jid]
            with progress_lock:
                for jid in stale:
                    progress_data.pop(jid, None)
            # Auto-unload model after idle timeout
            if (_ocr is not None
                    and _ocr_last_activity > 0
                    and (now - _ocr_last_activity) > MODEL_IDLE_TIMEOUT):
                logger.info("Model: 空闲超时，自动释放 OCR 模型")
                try:
                    unload_model()
                except Exception as e:
                    logger.warning(f"Model auto-unload error: {e}")
        except Exception as e:
            logger.warning(f"Cleanup error: {e}")
        _cleanup_stop.wait(CLEANUP_INTERVAL)


_cleanup_thread = threading.Thread(target=_cleanup_loop, daemon=True)
_cleanup_thread.start()
atexit.register(_cleanup_stop.set)


# ── Core OCR function ─────────────────────────────────────────────────────────
def ocr_page_image(img_array: np.ndarray) -> list:
    """
    Run PaddleOCR on one page image, returns list of {text, poly, score}.
    使用 paddleocr 3.x 新接口，兼容 predict() 返回格式。
    """
    ocr = get_ocr()
    try:
        results = ocr.predict(input=img_array)
    except Exception as e:
        logger.warning(f"OCR failed on page: {e}")
        return []

    items = []
    if not results:
        return items

    for res in results:
        # 3.x 返回对象，字段：rec_texts, rec_scores, rec_polys
        try:
            if hasattr(res, 'rec_texts'):
                texts  = res.rec_texts  or []
                scores = res.rec_scores or []
                polys  = res.rec_polys  or []
            elif isinstance(res, dict):
                texts  = res.get('rec_texts',  [])
                scores = res.get('rec_scores', [])
                polys  = res.get('rec_polys',  [])
            else:
                continue

            for i, (text, score) in enumerate(zip(texts, scores)):
                if float(score) < MIN_CONFIDENCE or not str(text).strip():
                    continue
                # poly: 取第 i 个多边形，格式为 [[x,y],...] 共4点
                poly = polys[i].tolist() if i < len(polys) else [[0,0],[0,0],[0,0],[0,0]]
                # 统一转为 [[x,y], ...] 列表格式（与旧接口一致）
                if hasattr(poly[0], '__iter__'):
                    poly = [list(p) for p in poly]
                items.append({
                    'text':  fix_ocr_text(str(text)),
                    'poly':  poly,
                    'score': float(score),
                })
        except Exception as e:
            logger.warning(f"OCR result parse error: {e}")
            continue

    return items


# ── Searchable PDF (PDF only, unchanged logic) ────────────────────────────────
def create_searchable_pdf(src_path: str, dst_path: str, task_id: str) -> dict:
    src_doc = fitz.open(src_path)
    if src_doc.is_encrypted:
        src_doc.close()
        raise RuntimeError("无法处理加密 PDF")
    total = len(src_doc)
    if total == 0:
        src_doc.close()
        raise RuntimeError("PDF 没有页面")

    new_doc = fitz.open()
    font    = fitz.Font("cjk")
    zoom    = OCR_DPI / PDF_PPI
    failed  = []

    update_progress(task_id, 0, total, 'processing', f'开始处理 {total} 页…')

    with ocr_lock:
        for i in range(total):
            if is_cancelled(task_id):
                src_doc.close(); new_doc.close()
                raise RuntimeError("已取消")

            update_progress(task_id, i + 1, total, 'processing', f'识别第 {i+1}/{total} 页…')

            page     = src_doc[i]
            rect     = page.rect
            rotation = page.rotation

            max_dim  = max(rect.width, rect.height)
            eff_zoom = zoom
            if max_dim * zoom > 4000:
                eff_zoom = max(150 / PDF_PPI, 4000 / max_dim / PDF_PPI)

            mat = fitz.Matrix(eff_zoom, eff_zoom)
            if rotation:
                mat = mat.prerotate(rotation)
            pix = page.get_pixmap(matrix=mat)
            img = pixmap_to_numpy(pix)

            try:
                ocr_data = ocr_page_image(img)
            except Exception as e:
                logger.error(f"Page {i+1} OCR error: {e}")
                ocr_data = []
                failed.append(i + 1)
            finally:
                del img
                del pix

            new_page = new_doc.new_page(width=rect.width, height=rect.height)
            pix2 = page.get_pixmap(matrix=mat)
            new_page.insert_image(rect, pixmap=pix2)
            del pix2

            sx = rect.width  / (rect.width  * eff_zoom)
            sy = rect.height / (rect.height * eff_zoom)
            row_tol = max(10, int(20 * OCR_DPI / 200))

            def sort_key(item):
                p = item['poly']
                y = min(pt[1] for pt in p)
                x = min(pt[0] for pt in p)
                return (int(y / row_tol), x)

            for item in sorted(ocr_data, key=sort_key):
                text = item['text']
                poly = item['poly']
                if not text or len(poly) < 4:
                    continue
                xs = [pt[0] for pt in poly]
                ys = [pt[1] for pt in poly]
                x0, y0 = min(xs) * sx, min(ys) * sy
                x1, y1 = max(xs) * sx, max(ys) * sy
                bw, bh = x1 - x0, y1 - y0
                if bw < 1 or bh < 1:
                    continue
                ref_fs = 10.0
                tl = font.text_length(text, fontsize=ref_fs)
                fs = (ref_fs * bw / tl) if tl > 0 else bh * 0.8
                fs = max(2.0, min(fs, 200.0))
                try:
                    tw = fitz.TextWriter(new_page.rect)
                    tw.append((x0, y1 - bh * 0.15), text, font=font, fontsize=fs)
                    tw.write_text(new_page, render_mode=3)
                except Exception:
                    pass

    src_doc.close()
    update_progress(task_id, total, total, 'saving', '保存 PDF…')
    new_doc.save(dst_path, garbage=4, deflate=True)
    new_doc.close()
    update_progress(task_id, total, total, 'done', '完成!')
    return {'total_pages': total, 'failed': failed}


# ── MD lines → docx helper ────────────────────────────────────────────────────
def _md_lines_to_docx(md_text: str, docx_path: Path) -> bool:
    """把 markdown 文字转存为 docx，先尝试 pandoc，失败则用 python-docx。"""
    # 先尝试 pandoc
    pandoc_exe = shutil.which('pandoc')
    if pandoc_exe:
        md_tmp = docx_path.with_suffix('.tmp.md')
        md_tmp.write_text(md_text, encoding='utf-8')
        try:
            subprocess.run(
                [pandoc_exe, str(md_tmp), '-o', str(docx_path),
                 '--from=markdown', '--to=docx'],
                timeout=120, check=True, capture_output=True,
            )
            md_tmp.unlink(missing_ok=True)
            if docx_path.exists():
                logger.info(f"Pandoc: DOCX created at {docx_path}")
                return True
        except Exception as e:
            logger.warning(f"Pandoc failed: {e}")
            md_tmp.unlink(missing_ok=True)

    # Fallback: python-docx
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt
        word = DocxDoc()
        for line in md_text.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith('## '):
                word.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                word.add_heading(line[4:], level=2)
            elif line.startswith('---'):
                word.add_paragraph('─' * 30)
            else:
                p = word.add_paragraph(line)
                p.style.font.size = Pt(11)
        word.save(str(docx_path))
        logger.info("python-docx: DOCX created as fallback")
        return True
    except Exception as e:
        logger.warning(f"python-docx fallback failed: {e}")
        return False


# ── OCR items → markdown lines ────────────────────────────────────────────────
def _ocr_items_to_md_lines(ocr_data: list, page_h: float) -> list[str]:
    """
    把单页 OCR 结果转为 markdown 行列表。
    使用 _group_lines_v2 进行行分组，解决手机拍摄畸变导致的排序乱问题。
    """
    if not ocr_data:
        return []

    lines     = _group_lines_v2(ocr_data)
    md_lines  = []

    for line_items in lines:
        text = ' '.join(it['text'] for it in line_items)
        if not text.strip():
            continue

        # 标题启发式检测：基于字符高度和位置
        avg_h = sum(
            (max(p[1] for p in it['poly']) - min(p[1] for p in it['poly']))
            for it in line_items
        ) / len(line_items)
        rel_y = min(p[1] for p in line_items[0]['poly']) / page_h if page_h > 0 else 0

        if avg_h > 30 and rel_y < 0.2 and len(text) < 60:
            md_lines.append(f"\n## {text}\n")
        elif avg_h > 22 and len(text) < 80:
            md_lines.append(f"\n### {text}\n")
        else:
            md_lines.append(text)

    return md_lines


# ── Markdown / DOCX — 三路分支 ────────────────────────────────────────────────
def create_markdown(src_path: str, out_dir: Path, task_id: str,
                    original_name: str = None,
                    clean_header: bool = False) -> dict:
    """
    三路分支：
      A. PDF 有文字层  → 直接提取，跳过 OCR
      B. PDF 纯图片    → OCR + _group_lines_v2
      C. JPG / PNG     → OCR + _group_lines_v2（单张图片）
    clean_header: 是否过滤页眉页脚（仅对 OCR 路径生效）
    """
    src      = Path(src_path)
    stem     = sanitize_name(Path(original_name or src_path).stem, 120)
    did      = f"{stem}_{task_id}"
    file_dir = out_dir / did
    file_dir.mkdir(parents=True, exist_ok=True)

    md_text     = ''
    total_pages = 0
    failed      = []

    # ── 路径 C：图片文件 ──────────────────────────────────────────────────────
    if src.suffix.lower() in {'.jpg', '.jpeg', '.png'}:
        logger.info(f"Image OCR mode: {src.name}")
        update_progress(task_id, 0, 1, 'processing', '加载图片…')

        try:
            img = load_image_as_numpy(str(src))
        except Exception as e:
            raise RuntimeError(f"图片加载失败: {e}")

        update_progress(task_id, 1, 1, 'processing', '识别图片文字…')

        # 记录图片高度，在 del img 之前
        img_h = img.shape[0] if hasattr(img, 'shape') else 1000

        with ocr_lock:
            if is_cancelled(task_id):
                shutil.rmtree(file_dir, ignore_errors=True)
                raise RuntimeError("已取消")
            ocr_data = ocr_page_image(img)

        del img

        if ocr_data:
            all_ys = [p[1] for it in ocr_data for p in it['poly']]
            page_h = max(all_ys) if all_ys else img_h
        else:
            page_h = img_h

        if ocr_data:
            if clean_header:
                ocr_data = _strip_header_footer(ocr_data, page_h)
            md_lines = _ocr_items_to_md_lines(ocr_data, page_h)
            md_text  = '\n'.join(md_lines)
        else:
            md_text = '*（未识别到文字）*'
            failed.append(1)

        total_pages = 1

    # ── 路径 A：PDF 有文字层 ──────────────────────────────────────────────────
    elif has_text_layer(src_path):
        logger.info(f"Text layer detected, skipping OCR: {src.name}")
        update_progress(task_id, 0, 1, 'processing', '检测到文字层，直接提取…')
        try:
            md_text = extract_text_layer(src_path)
        except Exception as e:
            raise RuntimeError(f"文字层提取失败: {e}")

        # 统计页数
        try:
            doc = fitz.open(src_path)
            total_pages = len(doc)
            doc.close()
        except Exception:
            total_pages = 1

        update_progress(task_id, 1, 1, 'processing', f'已提取 {total_pages} 页文字层')

    # ── 路径 B：PDF 纯图片，OCR ────────────────────────────────────────────────
    else:
        logger.info(f"Image-only PDF, OCR mode: {src.name}")
        doc   = fitz.open(src_path)
        if doc.is_encrypted:
            doc.close()
            raise RuntimeError("无法处理加密 PDF")
        total = len(doc)
        total_pages = total
        zoom  = OCR_DPI / PDF_PPI

        all_md_parts = []
        update_progress(task_id, 0, total, 'processing', f'开始 OCR {total} 页…')

        with ocr_lock:
            for i in range(total):
                if is_cancelled(task_id):
                    doc.close()
                    shutil.rmtree(file_dir, ignore_errors=True)
                    raise RuntimeError("已取消")

                update_progress(task_id, i + 1, total, 'processing', f'识别第 {i+1}/{total} 页…')

                page = doc[i]
                rect = page.rect
                mat  = fitz.Matrix(zoom, zoom)
                pix  = page.get_pixmap(matrix=mat)
                img  = pixmap_to_numpy(pix)
                del pix

                try:
                    ocr_data = ocr_page_image(img)
                except Exception as e:
                    logger.error(f"Page {i+1} OCR error: {e}")
                    ocr_data = []
                    failed.append(i + 1)
                finally:
                    del img

                if not ocr_data:
                    all_md_parts.append(f"\n\n---\n*第 {i+1} 页（无文字）*\n\n")
                    continue

                # 使用图像坐标系的页面高度（OCR 坐标基于渲染后的像素图）
                page_h_ocr = rect.height * zoom
                if clean_header:
                    ocr_data = _strip_header_footer(ocr_data, page_h_ocr)
                md_lines   = _ocr_items_to_md_lines(ocr_data, page_h_ocr)
                all_md_parts.append('\n'.join(md_lines))

        doc.close()
        md_text = '\n\n'.join(all_md_parts)
        full_memory_release()
    # ── 写出 md ───────────────────────────────────────────────────────────────
    update_progress(task_id, total_pages, total_pages, 'saving', '生成文档…')

    md_path = file_dir / f"{did}.md"
    md_path.write_text(md_text, encoding='utf-8')

    # ── 生成 docx ─────────────────────────────────────────────────────────────
    docx_path = file_dir / f"{did}.docx"
    docx_ok   = _md_lines_to_docx(md_text, docx_path)

    update_progress(task_id, total_pages, total_pages, 'done', '完成!')
    return {
        'total_pages': total_pages,
        'download_id': did,
        'has_docx':    docx_ok,
        'failed':      failed,
    }


# ── API routes ────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return send_file('static/index.html')


@app.route('/api/health')
def health():
    return jsonify({'status': 'ok', 'device': 'cpu', 'gpu_enabled': False})


@app.route('/api/csrf-token')
def csrf_token():
    return jsonify({'csrf_token': 'none'})


# ── Model management API ──────────────────────────────────────────────────────
@app.route('/api/model/status')
def model_status():
    return jsonify({
        'loaded':        _ocr is not None,
        'last_activity': _ocr_last_activity,
        'idle_timeout':  MODEL_IDLE_TIMEOUT,
    })


@app.route('/api/model/load', methods=['POST'])
def model_load():
    try:
        # 获取前端发送的 JSON 参数
        data = request.get_json(silent=True) or {}
        
        # 解析参数（如果没有传，则退化使用环境变量默认值）
        use_doc_ori  = data.get('use_doc_orientation', USE_DOC_ORIENTATION)
        use_unwarp   = data.get('use_doc_unwarping', USE_DOC_UNWARPING)
        use_text_ori = data.get('use_textline_orientation', USE_TEXTLINE_ORI)
        
        status = load_model(
            use_doc_ori=use_doc_ori,
            use_unwarp=use_unwarp,
            use_text_ori=use_text_ori
        )
        return jsonify({'status': status})
    except Exception as e:
        logger.error(f"Model load failed: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/model/unload', methods=['POST'])
def model_unload():
    unload_model()
    return jsonify({'status': 'unloaded'})


# ── Upload ────────────────────────────────────────────────────────────────────
@app.route('/api/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': '缺少文件'}), 400

    f    = request.files['file']
    mode = request.form.get('mode', 'pdf')
    if mode not in ('pdf', 'markdown'):
        mode = 'pdf'

    if not f.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = Path(f.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({'error': '仅支持 PDF、JPG、PNG 文件'}), 400

    # 图片文件不支持可搜索 PDF 模式
    if ext in {'.jpg', '.jpeg', '.png'} and mode == 'pdf':
        return jsonify({'error': '图片文件不支持可搜索 PDF 模式，请选择 Markdown/DOCX'}), 400

    # PDF 格式校验
    if ext == '.pdf' and not validate_pdf(f):
        return jsonify({'error': '文件不是有效的 PDF'}), 400

    # 引擎守卫
    if _ocr is None:
        return jsonify({'error': 'OCR 引擎未启动，请先在页面开启引擎'}), 503

    job_id   = str(uuid.uuid4())
    src_name = Path(f.filename).name
    src_path = UPLOAD_DIR / f"{job_id}_{src_name}"
    f.save(str(src_path))

    # Destination depends on mode
    if mode == 'pdf':
        out_stem = sanitize_name(Path(src_name).stem, 120)
        out_name = f"{out_stem}_{job_id}.pdf"
        output   = OUTPUT_DIR / out_name
    else:
        output = OUTPUT_DIR  # create_markdown 自己建子目录

    clean_header = request.form.get('clean_header', '0') == '1'

    now = time.time()
    with jobs_lock:
        jobs[job_id] = {
            'status':        'processing',
            'mode':          mode,
            'original_name': src_name,
            'created_at':    now,
            'updated_at':    now,
        }

    t = threading.Thread(
        target=_start_job,
        args=(job_id, mode, src_path, src_name, output, clean_header),
        daemon=True,
    )
    t.start()

    return jsonify({'job_id': job_id})


# ── Progress SSE ──────────────────────────────────────────────────────────────
@app.route('/api/progress/<job_id>')
def get_progress(job_id):
    def generate():
        import json
        last_sent = None
        timeout   = 600
        start     = time.time()
        while time.time() - start < timeout:
            with progress_lock:
                data = progress_data.get(job_id)
            with jobs_lock:
                job = jobs.get(job_id)

            if data and data != last_sent:
                last_sent = dict(data)
                payload   = dict(data)
                if job:
                    payload['job_status'] = job.get('status')
                    if job.get('status') == 'done':
                        payload['result'] = job.get('result')
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

            status = (job or {}).get('status', 'processing')
            if status in ('done', 'error', 'cancelled'):
                if job:
                    final = dict(data or {})
                    final['job_status'] = status
                    if job.get('result'):
                        final['result'] = job['result']
                    if job.get('error'):
                        final['error'] = job['error']
                    yield f"data: {json.dumps(final, ensure_ascii=False)}\n\n"
                break

            time.sleep(0.5)

    return Response(generate(), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route('/api/job/<job_id>')
def job_status(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Not found'}), 404
    with progress_lock:
        prog = progress_data.get(job_id, {})
    return jsonify({**job, 'progress': prog})


@app.route('/api/cancel/<job_id>', methods=['POST'])
def cancel_job(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            return jsonify({'error': 'Not found'}), 404
        if job['status'] != 'processing':
            return jsonify({'error': 'Not processing'}), 400
        job['status']     = 'cancelled'
        job['updated_at'] = time.time()
    set_cancelled(job_id)
    update_progress(job_id, 0, 0, 'cancelled', '已取消')
    return jsonify({'status': 'cancelled'})


# ── Download routes ────────────────────────────────────────────────────────────
@app.route('/api/download/pdf/<filename>')
def download_pdf(filename):
    try:
        p = validate_path(filename)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400
    if not p.exists():
        return jsonify({'error': 'Not found'}), 404
    orig = request.args.get('name', filename)
    safe = sanitize_name(Path(orig).stem)
    if not safe.lower().endswith('.pdf'):
        safe += '.pdf'
    return send_file(p, as_attachment=True, download_name=safe,
                     mimetype='application/pdf')


@app.route('/api/download/docx/<folder_name>')
def download_docx(folder_name):
    try:
        folder = validate_path(folder_name)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400
    docx = folder / f"{folder_name}.docx"
    if not docx.exists():
        return jsonify({'error': 'DOCX not found'}), 404
    orig = request.args.get('name', folder_name)
    safe = sanitize_name(Path(orig).stem) + '.docx'
    return send_file(docx, as_attachment=True, download_name=safe,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/download/markdown/<folder_name>')
def download_markdown(folder_name):
    import zipfile
    try:
        folder = validate_path(folder_name)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400
    if not folder.exists():
        return jsonify({'error': 'Not found'}), 404

    zip_path = OUTPUT_DIR / f"{folder_name}_{uuid.uuid4().hex[:8]}.zip"
    orig     = request.args.get('name', folder_name)
    safe     = sanitize_name(Path(orig).stem) + '.zip'

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for fp in folder.rglob('*'):
            if fp.is_file() and fp.suffix.lower() != '.docx':
                zf.write(fp, fp.relative_to(folder))

    @after_this_request
    def _del_zip(response):
        threading.Thread(
            target=lambda: (time.sleep(10), zip_path.unlink(missing_ok=True)),
            daemon=True,
        ).start()
        return response

    return send_file(zip_path, as_attachment=True, download_name=safe,
                     mimetype='application/zip')


@app.route('/api/view/markdown/<folder_name>')
def view_markdown(folder_name):
    try:
        folder = validate_path(folder_name)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400
    md = folder / f"{folder_name}.md"
    if not md.exists():
        return jsonify({'error': 'Not found'}), 404
    return jsonify({'content': md.read_text('utf-8'), 'filename': md.name})


@app.route('/api/delete/<path:file_id>', methods=['DELETE'])
def delete_export(file_id):
    try:
        p = validate_path(file_id)
    except ValueError:
        return jsonify({'error': 'Invalid path'}), 400
    try:
        if p.is_file():
            p.unlink()
        elif p.is_dir():
            shutil.rmtree(p)
        else:
            return jsonify({'status': 'not_found'}), 404
        return jsonify({'status': 'deleted'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Background job runner ─────────────────────────────────────────────────────
def _start_job(job_id: str, mode: str, src_path: Path,
               original_name: str, output_path_or_dir: Path,
               clean_header: bool = False):
    """Run OCR job in background thread."""
    try:
        if mode == 'pdf':
            result = create_searchable_pdf(str(src_path), str(output_path_or_dir), job_id)
            update_job(job_id, status='done', result={
                'mode':          'pdf',
                'download_id':   output_path_or_dir.name,
                'original_name': original_name,
                **result,
            })
        else:
            result = create_markdown(str(src_path), OUTPUT_DIR, job_id, original_name,
                                     clean_header=clean_header)
            update_job(job_id, status='done', result={
                'mode':          mode,
                'original_name': original_name,
                **result,
            })
    except RuntimeError as e:
        msg = str(e)
        if '已取消' in msg:
            update_job(job_id, status='cancelled')
        else:
            update_job(job_id, status='error', error=msg)
        update_progress(job_id, 0, 0, 'error', msg)
    except Exception as e:
        logger.exception(f"Job {job_id} unexpected error")
        update_job(job_id, status='error', error=str(e))
        update_progress(job_id, 0, 0, 'error', str(e))
    finally:
        try:
            Path(src_path).unlink(missing_ok=True)
        except Exception:
            pass


def update_job(job_id: str, **kwargs):
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            return
        cur = job.get('status')
        new = kwargs.get('status')
        if cur == 'cancelled' and new and new != 'cancelled':
            return
        job.update(kwargs)
        job['updated_at'] = time.time()


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True)
